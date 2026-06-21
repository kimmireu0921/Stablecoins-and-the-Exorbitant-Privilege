"""
make_docx.py — convert 0621_Final_Draft_for_Submission.md to a Word .docx
using python-docx (pandoc not available). Handles headings, paragraphs with
**bold**/*italic*, pipe tables, ![figures](){width}, code-fenced equations,
lists, and horizontal rules.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "0621_Final_Draft_for_Submission.md"
OUT = "0621_Final_Draft_for_Submission.docx"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
BLACK = RGBColor(0x00, 0x00, 0x00)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITAL_RE = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
FIG_RE  = re.compile(r"^!\[(.*)\]\((.+?)\)(?:\{.*\})?\s*$")


def inline_segments(text):
    """Return list of (text, bold, italic). Stray/triple asterisks stay literal."""
    text = text.replace("\\*", "\x00")
    parts, pos = [], 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False))
        parts.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False))

    final = []
    for seg, is_bold in parts:
        if is_bold:
            final.append((seg, True, False))
            continue
        ipos = 0
        for im in ITAL_RE.finditer(seg):
            if im.start() > ipos:
                final.append((seg[ipos:im.start()], False, False))
            final.append((im.group(1), False, True))
            ipos = im.end()
        if ipos < len(seg):
            final.append((seg[ipos:], False, False))
    return [(t.replace("\x00", "*"), b, i) for (t, b, i) in final if t]


def add_inline(p, text, size=12, font=BODY_FONT, color=None):
    for t, b, i in inline_segments(text):
        r = p.add_run(t)
        r.bold = b or None
        r.italic = i or None
        r.font.name = font
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color


def main():
    md = Path(SRC).read_text(encoding="utf-8").split("\n")
    doc = Document()

    # base style — font 12, single spacing, 1-inch margins (per submission rules)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(6)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(1)
        sec.top_margin = sec.bottom_margin = Inches(1)

    i, n = 0, len(md)
    title_done = False
    in_titleblock = False

    while i < n:
        line = md[i].rstrip("\n")
        stripped = line.strip()

        # ── code fence (equations) ──
        if stripped.startswith("```"):
            i += 1
            while i < n and not md[i].strip().startswith("```"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                add_inline(p, md[i].rstrip(), size=10.5, font=MONO_FONT)
                i += 1
            i += 1
            continue

        # ── blank ──
        if stripped == "":
            i += 1
            continue

        # ── horizontal rule ──
        if stripped == "---":
            in_titleblock = False
            i += 1
            continue

        # ── headings ──
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            htext = stripped.lstrip("#").strip()
            if level == 1 and not title_done:
                # two-line centered, bold, underlined, black title (split at first colon)
                if ": " in htext:
                    main, sub = htext.split(": ", 1)
                    main = main + ":"
                else:
                    main, sub = htext, None
                for txt, sz in [(main, 16), (sub, 14)]:
                    if not txt:
                        continue
                    tp = doc.add_paragraph()
                    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tp.paragraph_format.space_after = Pt(2)
                    r = tp.add_run(txt)
                    r.bold = True
                    r.underline = True
                    r.font.name = BODY_FONT
                    r.font.size = Pt(sz)
                    r.font.color.rgb = BLACK
                title_done = True
                in_titleblock = True   # author/affiliation lines follow
            else:
                # manual bold-black heading (avoids Word's default blue theme color)
                hp = doc.add_paragraph()
                hp.paragraph_format.space_before = Pt(10)
                hp.paragraph_format.space_after = Pt(4)
                sz = {2: 14, 3: 12}.get(level, 12)
                for seg, b, it in inline_segments(htext):
                    r = hp.add_run(seg)
                    r.bold = True
                    r.italic = it or None
                    r.font.name = BODY_FONT
                    r.font.size = Pt(sz)
                    r.font.color.rgb = BLACK
            i += 1
            continue

        # ── figure ──
        mfig = FIG_RE.match(stripped)
        if mfig:
            caption, path = mfig.group(1), mfig.group(2)
            if Path(path).exists():
                doc.add_picture(path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(cap, caption, size=9)
            i += 1
            continue

        # ── table (consecutive lines starting with |) ──
        if stripped.startswith("|"):
            rows = []
            while i < n and md[i].strip().startswith("|"):
                rows.append([c.strip() for c in md[i].strip().strip("|").split("|")])
                i += 1
            # drop separator row(s) like |---|---|
            data = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if not data:
                continue
            ncol = max(len(r) for r in data)
            tbl = doc.add_table(rows=len(data), cols=ncol)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(data):
                for ci in range(ncol):
                    cell = tbl.cell(ri, ci)
                    cell.paragraphs[0].text = ""
                    txt = row[ci] if ci < len(row) else ""
                    add_inline(cell.paragraphs[0], txt, size=9.5)
                    if ri == 0:
                        for r in cell.paragraphs[0].runs:
                            r.bold = True
            doc.add_paragraph()
            continue

        # ── bullet / numbered list ──
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # ── normal paragraph (title block lines centered, body justified) ──
        p = doc.add_paragraph()
        if in_titleblock:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print(f"Saved {OUT}")
    # validation
    d2 = Document(OUT)
    ntab = len(d2.tables)
    nimg = sum(1 for s in d2.inline_shapes)
    print(f"  paragraphs={len(d2.paragraphs)}  tables={ntab}  images={nimg}")


if __name__ == "__main__":
    main()
